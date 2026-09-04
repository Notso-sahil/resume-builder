from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

from src.schemas.models import ResumeProjectPortfolio, CandidateProfile


def add_horizontal_rule(paragraph):
    """Adds a crisp bottom border line beneath a section heading."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     r'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="333333"/>'
                     r'</w:pBdr>')
    pPr.append(pBdr)


def render_docx(
    portfolio: ResumeProjectPortfolio,
    candidate: Optional[CandidateProfile],
    output_path: str | Path,
) -> Path:
    """
    Renders the synthesized resume into a dense, high-impact Microsoft Word document (.docx)
    that completely fills exactly ONE single page with rich engineering depth.

    Order:
      1. Header (Name, Title, Contact)
      2. Professional Summary (tailored, blending background + JD)
      3. Experience (AI Research Intern at IFSO, Delhi Police)
      4. Technical Skills (categorized, placed before projects)
      5. Technical Projects (3 archetypes, rich Google XYZ bullets)
      6. Education (placed at bottom)
    """
    doc = Document()
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    # 1. Page Margins calibrated to completely fill exactly 1 full letter page
    for section in doc.sections:
        section.top_margin = Inches(0.33)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    # 2. Candidate Header
    name = candidate.full_name if candidate else "CANDIDATE NAME"
    title = candidate.title if candidate and candidate.title else (
        portfolio.jd_analysis.role_title if portfolio.jd_analysis else "AI Systems Engineer"
    )
    phone = candidate.phone if candidate else "+91 8700122453"
    email = candidate.email if candidate else "candidate@example.com"
    linkedin = candidate.linkedin if candidate else None
    github = candidate.github if candidate else None

    # Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1)
    run_name = p_name.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(15.5)
    run_name.font.name = "Calibri"

    # Title / Headline
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(1)
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(9.5)
    run_title.font.name = "Calibri"
    run_title.font.color.rgb = RGBColor(70, 70, 70)

    # Contact line
    contacts = [phone, email]
    if linkedin:
        contacts.append(linkedin)
    if github:
        contacts.append(github)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(2.5)
    run_contact = p_contact.add_run("   |   ".join(contacts))
    run_contact.font.size = Pt(8.5)
    run_contact.font.name = "Calibri"

    # Section Header Helper
    def add_section_header(title_text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(title_text.upper())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(20, 20, 20)
        add_horizontal_rule(p)
        return p

    # 3. Professional Summary (dynamically tailored)
    summary_text = (
        (candidate.tailored_summary if candidate else None)
        or portfolio.tailored_summary
        or (candidate.professional_objective if candidate else None)
    )
    if summary_text:
        add_section_header("Professional Summary")
        p_obj = doc.add_paragraph()
        p_obj.paragraph_format.space_before = Pt(1)
        p_obj.paragraph_format.space_after = Pt(3)
        p_obj.paragraph_format.line_spacing = 1.1
        run_obj = p_obj.add_run(summary_text)
        run_obj.font.size = Pt(8.5)
        run_obj.font.name = "Calibri"

    # 4. Work Experience
    if candidate and candidate.experience:
        add_section_header("Work Experience")
        for exp in candidate.experience:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(2)
            p_exp.paragraph_format.space_after = Pt(1)

            run_role = p_exp.add_run(exp.role)
            run_role.bold = True
            run_role.font.size = Pt(9.5)
            run_role.font.name = "Calibri"

            run_org = p_exp.add_run(f"  |  {exp.organization}")
            run_org.italic = True
            run_org.font.size = Pt(9)
            run_org.font.name = "Calibri"
            run_org.font.color.rgb = RGBColor(50, 50, 50)

            if exp.location:
                run_loc = p_exp.add_run(f"  ·  {exp.location}")
                run_loc.font.size = Pt(8.5)
                run_loc.font.name = "Calibri"
                run_loc.font.color.rgb = RGBColor(80, 80, 80)

            if exp.period:
                run_per = p_exp.add_run(f"    ({exp.period})")
                run_per.font.size = Pt(8.5)
                run_per.font.name = "Calibri"
                run_per.font.color.rgb = RGBColor(80, 80, 80)

            for b in exp.bullets:
                p_b = doc.add_paragraph(style="List Bullet")
                p_b.paragraph_format.space_before = Pt(0.5)
                p_b.paragraph_format.space_after = Pt(1)
                p_b.paragraph_format.line_spacing = 1.1
                run_b = p_b.add_run(b)
                run_b.font.size = Pt(8.5)
                run_b.font.name = "Calibri"

    # 5. Technical Skills (SWAPPED - placed before Technical Projects)
    if portfolio.jd_analysis:
        jd = portfolio.jd_analysis
        add_section_header("Technical Skills")

        def add_skill_line(label: str, items: list):
            if items:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0.5)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05
                r_lbl = p.add_run(f"{label}: ")
                r_lbl.bold = True
                r_lbl.font.size = Pt(8.5)
                r_lbl.font.name = "Calibri"
                r_val = p.add_run(", ".join(items))
                r_val.font.size = Pt(8.5)
                r_val.font.name = "Calibri"

        add_skill_line("Languages & Core", jd.primary_languages)
        add_skill_line("Inference Engines & LLMs", jd.frameworks)
        add_skill_line("Distributed Systems & Cloud", jd.databases_and_storage + jd.infrastructure_and_cloud)
        opt_skills = [
            "Model inference optimization", "GPU benchmarking", "Speculative decoding",
            "KV-cache strategy", "Quantization (AWQ/FP8)", "Continuous batching",
            "Tensor & pipeline parallelism", "Memory bandwidth", "API usage for LLMs"
        ]
        add_skill_line("Optimization & Profiling", opt_skills)

    # 6. Technical Projects (3 projects, high-impact XYZ bullets)
    add_section_header("Technical Projects")
    for proj in portfolio.projects:
        p_proj = doc.add_paragraph()
        p_proj.paragraph_format.space_before = Pt(2)
        p_proj.paragraph_format.space_after = Pt(0.5)

        run_pname = p_proj.add_run(proj.project_title)
        run_pname.bold = True
        run_pname.font.size = Pt(9.5)
        run_pname.font.name = "Calibri"

        stack_text = "  |  " + ", ".join(proj.tech_stack[:6])
        run_pstack = p_proj.add_run(stack_text)
        run_pstack.italic = True
        run_pstack.font.size = Pt(8.5)
        run_pstack.font.name = "Calibri"
        run_pstack.font.color.rgb = RGBColor(60, 60, 60)

        # Render 3-4 dense bullets per project to fill page with substance
        bullets_to_render = proj.xyz_bullets[:4]
        for bullet in bullets_to_render:
            p_bullet = doc.add_paragraph(style="List Bullet")
            p_bullet.paragraph_format.space_before = Pt(0)
            p_bullet.paragraph_format.space_after = Pt(0.5)
            p_bullet.paragraph_format.line_spacing = 1.05
            run_b = p_bullet.add_run(bullet)
            run_b.font.size = Pt(8.5)
            run_b.font.name = "Calibri"

    # 7. Education (SWAPPED - placed cleanly at bottom)
    if candidate and candidate.education:
        add_section_header("Education")
        for edu in candidate.education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(1)
            p_edu.paragraph_format.space_after = Pt(1)

            run_deg = p_edu.add_run(edu.degree)
            run_deg.bold = True
            run_deg.font.size = Pt(9)
            run_deg.font.name = "Calibri"

            if edu.year_range:
                run_yr = p_edu.add_run(f"    ({edu.year_range})")
                run_yr.font.size = Pt(8.5)
                run_yr.font.name = "Calibri"
                run_yr.font.color.rgb = RGBColor(80, 80, 80)

            p_inst = doc.add_paragraph()
            p_inst.paragraph_format.space_before = Pt(0)
            p_inst.paragraph_format.space_after = Pt(2)
            inst_text = edu.institution
            if edu.details:
                inst_text += f"   |   {edu.details}"
            run_inst = p_inst.add_run(inst_text)
            run_inst.font.size = Pt(8.5)
            run_inst.font.name = "Calibri"
            run_inst.font.color.rgb = RGBColor(70, 70, 70)

    doc.save(str(out))
    return out
